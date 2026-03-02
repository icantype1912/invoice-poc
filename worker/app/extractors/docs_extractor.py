import io
import logging
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(docx_data: bytes) -> str:
    """
    Extract text from DOCX file.
    Args:
        docx_data: Raw DOCX bytes
    Returns:
        Extracted text string
    Raises:
        Exception: If extraction fails
    """
    try:
        logger.debug(f"Opening DOCX ({len(docx_data)} bytes)")

        text_parts = []

        document = Document(io.BytesIO(docx_data))
        logger.debug(f"DOCX has {len(document.paragraphs)} paragraphs")

        for idx, para in enumerate(document.paragraphs, 1):
            text = para.text.strip()
            if text:
                text_parts.append(text)
                logger.debug(f"Paragraph {idx}: {len(text)} characters")

        full_text = "\n\n".join(text_parts)
        logger.debug(f"Total extracted: {len(full_text)} characters")

        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}", exc_info=True)
        raise Exception(f"DOCX extraction failed: {str(e)}")
